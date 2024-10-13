import re
import io
import struct
import datetime

from tqdm import tqdm
from docx import Document
from loguru import logger
from flask import Flask, Response, abort, render_template, flash, request, redirect, url_for, send_file

from forms import TranslateForm, RAGForm
from gpt_api_class import GPT_API, RAG

app = Flask(__name__)
gpt_api = GPT_API()
gpt_rag = RAG()

app.config['WTF_CSRF_ENABLED'] = False
app.secret_key = 'super secret key'


def has_cyrillic(text):
    return bool(re.search('[а-яА-Я]', text))


@app.route('/')
def main():  
     return render_template('main.html')


@app.route('/translate', methods=["POST", "GET"])
def translate():   
    form = TranslateForm()

    if request.method == 'GET':
        return render_template('translate.html', form=form)

    if request.method == 'POST':
        if form.validate_on_submit():
            data = io.BytesIO(form.file.data.stream.read()) 
            doc = Document(data)

            total_cost = 0
            for paragraph in tqdm(doc.paragraphs):
                for run in paragraph.runs:
                    current_text = run.text.strip()
                    if len(current_text) > 0 and has_cyrillic(current_text):
                        translated_text, cost = gpt_api.invoke(text=current_text)
                        total_cost += cost
                        run.text = translated_text

            for table in tqdm(doc.tables):
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                current_text = run.text.strip()
                                if len(current_text) > 0 and has_cyrillic(current_text):
                                    translated_text, cost = gpt_api.invoke(text=current_text)
                                    total_cost += cost
                                    run.text = translated_text
            logger.info(total_cost)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return send_file(
                buffer,
                as_attachment=True,
                download_name='Test.docx',
                mimetype='application/msword'
            )

        flash("Incorrect form")
        return render_template('translate.html', form=form)
    return redirect(url_for('translate'), form=form)


@app.route('/rag', methods=["POST", "GET"])
def rag():
    form = RAGForm()

    if request.method == 'GET':
        return render_template('rag.html', form=form)

    if request.method == 'POST':
        if form.validate_on_submit():
            logger.info(f"Запрос: {form.query.data}")
            answer = gpt_rag.execute_query(form.query.data)
            print(answer)
            return render_template('rag.html', form=form, answer=answer) 
        flash("Incorrect form")

    return render_template('rag.html', form=form)

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=8000, debug=True)